import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Reporte de Examen: Solucionador N-Puzzle con IDA*', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1
    doc.add_heading('1. Lectura de Archivo y Salida de Consola', level=1)
    doc.add_paragraph(
        "El programa principal (ahora llamado principal.py) lee el archivo de entrada .txt que "
        "incluye la dimensión del tablero, la matriz inicial y la matriz meta. La salida se imprime "
        "estrictamente en consola con el formato solicitado de movimientos separados por comas (ej. U,D,L,R)."
    )

    # Section 2
    doc.add_heading('2. Pseudocódigo y Lógica IDA*', level=1)
    doc.add_paragraph(
        "El motor del algoritmo está en ida_estrella.py. Se respetó la arquitectura estricta del IDA*:\n"
        "• Se implementó un control Tabú por rama de profundidad para evitar ciclos absolutos.\n"
        "• Se integró una capa inicial BFS (Anchura) para crear una Tabla de Transposiciones que "
        "impide re-visitar los estados fundacionales (mejorando el rendimiento).\n"
        "• El control de profundidad se maneja mediante un 'threshold' limitante."
    )

    # Section 3
    doc.add_heading('3. Función de Evaluación y Heurísticas Activas', level=1)
    doc.add_paragraph(
        "La evaluación (f(n) = g(n) + h(n)) se lleva a cabo en heuristicas.py utilizando arreglos "
        "planos (1D) para maximizar la velocidad computacional en Python. Se implementaron 6 componentes:"
    )
    doc.add_paragraph("   1. Distancia Manhattan (MD)", style='List Bullet')
    doc.add_paragraph("   2. Conflicto Lineal (LC)", style='List Bullet')
    doc.add_paragraph("   3. Fichas en Esquinas (Corner Tiles - CT)", style='List Bullet')
    doc.add_paragraph("   4. Último Movimiento (Last Move - LM)", style='List Bullet')
    doc.add_paragraph("   5. Distancia Caminando (Walking Distance - WD)", style='List Bullet')
    doc.add_paragraph("   6. Distancia de Inversión (Inversion Distance - IDR)", style='List Bullet')
    
    doc.add_paragraph(
        "Admisibilidad (Corrección Matemática): Para garantizar siempre hallar el camino más corto, "
        "no se suman las heurísticas ciegamente (lo que sobreestimaría y rompería el IDA*). En su lugar, "
        "se calcula conservadoramente el escenario máximo: f(n) = max(MD+LC+CT, MD+LC+LM, MD+IDR, WD)."
    )

    # Section 4
    doc.add_heading('4. Implementación en Python', level=1)
    doc.add_paragraph(
        "Se desarrolló enteramente en Python 3 utilizando paradigmas de optimización extrema (evitando "
        "clonaciones profundas 2D). Los archivos del código base fueron traducidos al español para mayor "
        "comodidad en la lectura técnica."
    )
    
    # Section 5
    doc.add_heading('5. Análisis Empírico Empaquetado (Dashboard y Resultados)', level=1)
    doc.add_paragraph(
        "Se diseñó el módulo visualizador.py con la biblioteca Seaborn para analizar las 1,800 instancias de prueba del examen. La imagen generada a continuación corresponde al Dashboard estadístico de los datos actualizados tras la corrección heurística."
    )

    # Add Image
    img_path = '/Users/nglmike/Downloads/IA/n_puzzle_solver/results/stats_plots/resultados_completos_dashboard.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        p = doc.add_paragraph("Figura 1: Dashboard Multivariable de Rendimiento y Costos")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("(No se encontró la imagen del dashboard en la ruta esperada).")
        
    # Save Report
    report_path = '/Users/nglmike/Downloads/IA/n_puzzle_solver/Reporte_Examen_IDA.docx'
    doc.save(report_path)
    print(f"Reporte Word generado satisfactoriamente en: {report_path}")

if __name__ == "__main__":
    main()
